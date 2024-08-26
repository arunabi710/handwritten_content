import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from io import BytesIO
from docx import Document
import tempfile

# Azure Form Recognizer credentials
endpoint = "https://arun-document-ai.cognitiveservices.azure.com/"
key = "d3abb1fb970e41d8b7f3330e202f342a"

def format_bounding_box(bounding_box):
    if not bounding_box:
        return "N/A"
    return ", ".join(["[{}, {}]".format(p.x, p.y) for p in bounding_box])

def analyze_read(file_content):
    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )

    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(file_content)
        temp_file_path = temp_file.name

    with open(temp_file_path, "rb") as file:
        poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file
        )
        result = poller.result()

    doc = Document()
    doc.add_heading('Analysis Results', 0)

    # Concatenate the content of the document
    full_content = " ".join([line.content for page in result.pages for line in page.lines])
    doc.add_paragraph("Document contains content: {}".format(full_content))
    
    has_handwritten_content = any(style.is_handwritten for style in result.styles)
    if has_handwritten_content:
        doc.add_paragraph("Document contains handwritten content")
    
    return doc

st.title("Document Intelligence Tool")

uploaded_file = st.file_uploader("Upload a PDF or image file", type=["pdf", "jpg", "jpeg", "png"])

if uploaded_file is not None:
    file_content = uploaded_file.read()
    doc = analyze_read(file_content)
    
    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    st.download_button(
        label="Download Analysis Result",
        data=output,
        file_name="analysis_result.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
